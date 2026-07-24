"""
Microsoft Word (DOCX) report generator using python-docx.

Produces a corporate-style Word document with:
  • Cover page
  • Executive summary table
  • Quality overview
  • Issues summary table
  • Detailed issues with recommendations
  • Top files / top rules sections
  • Technical debt summary
"""

from __future__ import annotations

import io
import logging
from typing import TYPE_CHECKING, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from models.issue import IssueType
from models.metrics import rating_color, rating_letter
from models.project import ReportData
from reports.base_report import BaseReport
from services.chart_service import ChartService
from services.recommendation_engine import RecommendationEngine

if TYPE_CHECKING:
    pass

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Colour constants (RGB tuples)
# ---------------------------------------------------------------------------

C_PRIMARY   = RGBColor(0x1A, 0x23, 0x7E)
C_SECONDARY = RGBColor(0x28, 0x35, 0x93)
C_ACCENT    = RGBColor(0x42, 0xA5, 0xF5)
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
C_SUCCESS   = RGBColor(0x4C, 0xAF, 0x50)
C_DANGER    = RGBColor(0xF4, 0x43, 0x36)
C_WARNING   = RGBColor(0xFF, 0x98, 0x00)
C_DARK      = RGBColor(0x21, 0x21, 0x21)
C_MEDIUM    = RGBColor(0x54, 0x6E, 0x7A)
C_LIGHT_BG  = RGBColor(0xF0, 0xF4, 0xFF)

RATING_COLORS = {
    "A": RGBColor(0x4C, 0xAF, 0x50),
    "B": RGBColor(0x8B, 0xC3, 0x4A),
    "C": RGBColor(0xFF, 0x98, 0x00),
    "D": RGBColor(0xF4, 0x43, 0x36),
    "E": RGBColor(0xD3, 0x2F, 0x2F),
}

SEV_COLORS = {
    "BLOCKER":  RGBColor(0xD3, 0x2F, 0x2F),
    "CRITICAL": RGBColor(0xF4, 0x43, 0x36),
    "MAJOR":    RGBColor(0xFF, 0x98, 0x00),
    "MINOR":    RGBColor(0xFF, 0xC1, 0x07),
    "INFO":     RGBColor(0x21, 0x96, 0xF3),
}


# ---------------------------------------------------------------------------
# Helper utilities
# ---------------------------------------------------------------------------

def _rgb_hex(r: int, g: int, b: int) -> str:
    return f"{r:02X}{g:02X}{b:02X}"


def _set_cell_bg(cell, rgb: RGBColor) -> None:
    """Set cell background colour via XML (python-docx doesn't expose this directly)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    # RGBColor is a bytes tuple (r, g, b) — no .red/.green/.blue attributes
    hex_ = _rgb_hex(rgb[0], rgb[1], rgb[2])
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_)
    tcPr.append(shd)


def _set_para_spacing(para, before: int = 0, after: int = 4) -> None:
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"),  str(after  * 20))
    pPr.append(spacing)


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal rule."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A237E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _set_para_spacing(para, after=2)


# ---------------------------------------------------------------------------
# DOCX Report
# ---------------------------------------------------------------------------

class DocxReport(BaseReport):
    """Generates a professional Word document report."""

    def __init__(self, output_dir: str = "output") -> None:
        super().__init__(output_dir)
        self._chart_svc  = ChartService()
        self._rec_engine = RecommendationEngine()

    def generate(self, data: ReportData) -> str:
        path = self.output_path("Report.docx")
        logger.info("Generating DOCX report → %s", path)

        doc = Document()
        self._configure_document(doc)

        self._cover_page(doc, data)
        self._executive_summary(doc, data)
        self._quality_overview(doc, data)
        self._charts_section(doc, data)
        self._issues_summary(doc, data)
        self._detailed_issues(doc, data)
        self._top_files(doc, data)
        self._top_rules(doc, data)
        self._technical_debt(doc, data)

        doc.save(path)
        logger.info("DOCX report saved")
        return path

    # ------------------------------------------------------------------
    # Document configuration
    # ------------------------------------------------------------------

    def _configure_document(self, doc: Document) -> None:
        """Set default font and margins."""
        section = doc.sections[0]
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

        # Default style
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------

    def _cover_page(self, doc: Document, data: ReportData) -> None:
        doc.add_paragraph()
        doc.add_paragraph()

        title = doc.add_heading(data.report_title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.font.color.rgb = C_PRIMARY
        run.font.size      = Pt(26)
        _set_para_spacing(title, before=20, after=12)

        subtitle = doc.add_paragraph(data.project_name)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.runs[0]
        run.font.size      = Pt(18)
        run.font.color.rgb = C_SECONDARY
        run.bold           = True
        _set_para_spacing(subtitle, after=20)

        _add_horizontal_rule(doc)

        meta = [
            ("Prepared for", data.company_name),
            ("Project Key", data.project_key),
            ("Generated",   self.format_datetime(data.generated_at)),
            ("Quality Gate", data.metrics.quality_gate_status),
            ("Total Issues", str(len(data.issues))),
            ("Lines of Code", f"{data.metrics.ncloc:,}"),
        ]

        tbl = doc.add_table(rows=len(meta), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(meta):
            row = tbl.rows[i]
            lc = row.cells[0]
            vc = row.cells[1]
            lc.text = label
            vc.text = value
            lc.paragraphs[0].runs[0].bold = True
            lc.paragraphs[0].runs[0].font.color.rgb = C_PRIMARY
            _set_cell_bg(lc, C_LIGHT_BG)

        doc.add_paragraph()
        disclaimer = doc.add_paragraph("STRICTLY CONFIDENTIAL – For authorized personnel only")
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.runs[0]
        run.font.size = Pt(8)
        run.italic    = True
        run.font.color.rgb = C_MEDIUM

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Executive Summary
    # ------------------------------------------------------------------

    def _executive_summary(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "1. Executive Summary")
        m = data.metrics

        # Gate
        gate_ok    = m.quality_gate_status == "OK"
        gate_label = "PASSED ✓" if gate_ok else "FAILED ✗"
        gate_para  = doc.add_paragraph()
        run        = gate_para.add_run(f"Quality Gate: {gate_label}")
        run.bold   = True
        run.font.size = Pt(12)
        run.font.color.rgb = C_SUCCESS if gate_ok else C_DANGER

        doc.add_paragraph(
            f"This report presents a comprehensive code quality analysis of the {data.project_name} "
            f"project ({data.project_key}), generated on {self.format_date(data.generated_at)}. "
            f"The analysis covers {m.ncloc:,} lines of code across {m.files} files and identifies "
            f"{len(data.issues)} quality issues requiring attention."
        )

        # Summary table
        rows_data = [
            ["Metric", "Value", "Rating"],
            ["Lines of Code", f"{m.ncloc:,}", "–"],
            ["Reliability (Bugs)", f"{m.bugs}", m.reliability_rating_letter],
            ["Security (Vulns)", f"{m.vulnerabilities}", m.security_rating_letter],
            ["Maintainability (Smells)", f"{m.code_smells}", m.sqale_rating_letter],
            ["Technical Debt", m.technical_debt_display, "–"],
            ["Coverage", f"{m.coverage:.1f}%", "–"],
            ["Duplicated Lines", f"{m.duplicated_lines_density:.1f}%", "–"],
            ["Open Issues", str(m.open_issues), "–"],
            ["Accepted/Resolved Issues", str(m.resolved_issues), "–"],
        ]

        tbl = doc.add_table(rows=len(rows_data), cols=3)
        tbl.style = "Table Grid"
        for r_idx, row_vals in enumerate(rows_data):
            for c_idx, val in enumerate(row_vals):
                cell = tbl.rows[r_idx].cells[c_idx]
                cell.text = val
                run  = cell.paragraphs[0].runs[0]
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = C_WHITE
                    _set_cell_bg(cell, C_PRIMARY)
                else:
                    if c_idx == 2 and val in RATING_COLORS:
                        run.bold = True
                        run.font.color.rgb = C_WHITE
                        _set_cell_bg(cell, RATING_COLORS[val])
                    elif r_idx % 2 == 0:
                        _set_cell_bg(cell, C_LIGHT_BG)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Quality Overview
    # ------------------------------------------------------------------

    def _quality_overview(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "2. Quality Overview")
        m = data.metrics

        cards = [
            ("Reliability",     f"{m.bugs} Bug(s)",                m.reliability_rating_letter),
            ("Maintainability", f"{m.code_smells} Code Smell(s)",  m.sqale_rating_letter),
            ("Security",        f"{m.vulnerabilities} Vuln(s)",    m.security_rating_letter),
            ("Coverage",        f"{m.coverage:.1f}%",              "–"),
            ("Duplications",    f"{m.duplicated_lines_density:.1f}%","–"),
            ("Technical Debt",  m.technical_debt_display,           "–"),
        ]

        tbl = doc.add_table(rows=2, cols=3)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        flat_cards = cards[:6]
        for i, (cat, val, rating) in enumerate(flat_cards):
            row_idx = i // 3
            col_idx = i % 3
            cell = tbl.rows[row_idx].cells[col_idx]
            cell.text = ""
            p1 = cell.paragraphs[0]
            r1 = p1.add_run(cat.upper())
            r1.bold = True
            r1.font.color.rgb = C_WHITE
            r1.font.size = Pt(9)
            _set_cell_bg(cell, C_PRIMARY)

            p2 = cell.add_paragraph(val)
            p2.runs[0].bold = True
            p2.runs[0].font.size = Pt(14)
            p2.runs[0].font.color.rgb = C_ACCENT

            p3 = cell.add_paragraph(f"Rating: {rating}")
            p3.runs[0].font.size = Pt(8)
            p3.runs[0].font.color.rgb = C_MEDIUM

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Charts
    # ------------------------------------------------------------------

    def _charts_section(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "3. Charts & Visualisations")
        cs = self._chart_svc

        chart_pairs = [
            ("Severity Distribution",    cs.severity_pie(data)),
            ("Issue Type Distribution",  cs.type_pie(data)),
            ("Issues by Severity",       cs.severity_bar(data)),
            ("Top Files",                cs.issues_per_file_bar(data)),
            ("Top Rules",                cs.top_rules_bar(data)),
        ]

        for title, chart_bytes in chart_pairs:
            doc.add_paragraph(title).runs[0].bold = True
            buf = io.BytesIO(chart_bytes)
            doc.add_picture(buf, width=Cm(14))
            doc.add_paragraph()

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Issues Summary
    # ------------------------------------------------------------------

    def _issues_summary(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "4. Issues Summary")
        doc.add_paragraph(f"Total: {len(data.issues)} issues sorted by severity.")

        headers = ["#", "Severity", "Type", "Rule", "File", "Line", "Status"]
        tbl = doc.add_table(rows=1 + len(data.issues), cols=len(headers))
        tbl.style = "Table Grid"

        # Header
        for c_idx, h in enumerate(headers):
            cell = tbl.rows[0].cells[c_idx]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
            _set_cell_bg(cell, C_PRIMARY)

        for r_idx, issue in enumerate(data.issues_sorted_by_severity, start=1):
            row = tbl.rows[r_idx]
            vals = [
                str(r_idx),
                issue.severity.value,
                issue.issue_type.label,
                issue.rule.split(":")[-1][:20],
                self.truncate(issue.file_name, 30),
                str(issue.display_line or "–"),
                issue.issue_status.value,
            ]
            for c_idx, val in enumerate(vals):
                cell = row.cells[c_idx]
                cell.text = val
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if c_idx == 1:
                    color = SEV_COLORS.get(issue.severity.value)
                    if color:
                        run.bold = True
                        run.font.color.rgb = C_WHITE
                        _set_cell_bg(cell, color)
                elif r_idx % 2 == 0:
                    _set_cell_bg(cell, C_LIGHT_BG)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Detailed Issues
    # ------------------------------------------------------------------

    def _detailed_issues(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "5. Detailed Issues")

        for idx, issue in enumerate(data.issues_sorted_by_severity, start=1):
            rec  = self._rec_engine.recommend(issue, data.sonar_url)
            rule = data.rules.get(issue.rule, {})

            # Issue heading
            h = doc.add_heading(f"Issue #{idx} – {issue.rule}", level=3)
            h.runs[0].font.color.rgb = C_PRIMARY
            _set_para_spacing(h, before=12, after=4)

            # Meta table
            meta = doc.add_table(rows=3, cols=4)
            meta.style = "Table Grid"
            meta_data = [
                ("Severity", issue.severity.value, "Type", issue.issue_type.label),
                ("File", self.truncate(issue.file_path, 50), "Line", str(issue.display_line or "–")),
                ("Status", issue.issue_status.value, "Effort", issue.effort),
            ]
            for r_idx, (k1, v1, k2, v2) in enumerate(meta_data):
                row = meta.rows[r_idx]
                for c_idx, val in enumerate([k1, v1, k2, v2]):
                    cell = row.cells[c_idx]
                    cell.text = val
                    run = cell.paragraphs[0].runs[0]
                    run.font.size = Pt(8)
                    if c_idx % 2 == 0:
                        run.bold = True
                        _set_cell_bg(cell, C_LIGHT_BG)

            doc.add_paragraph()

            # Problem
            self._sub_heading(doc, "Problem")
            doc.add_paragraph(issue.message)

            if issue.code_snippet:
                self._sub_heading(doc, "Code")
                snippet = doc.add_paragraph()
                snippet.paragraph_format.space_after = Pt(8)
                for line in issue.code_snippet.splitlines():
                    run = snippet.add_run(line + "\n")
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)

            # Why
            rule_desc = self._strip_html(rule.get("mdDesc", rule.get("htmlDesc", "")))
            if rule_desc:
                self._sub_heading(doc, "Why Is This an Issue?")
                doc.add_paragraph(rule_desc[:600])

            # Developer Recommendation
            self._sub_heading(doc, "Developer Recommendation")
            doc.add_paragraph(rec.developer_recommendation)

            # Business Impact
            self._sub_heading(doc, "Business Impact")
            doc.add_paragraph(rec.business_impact)

            # Priority / Difficulty / Time
            pdt_tbl = doc.add_table(rows=1, cols=3)
            pdt_tbl.style = "Table Grid"
            for c_idx, (label, val) in enumerate([
                ("Priority", rec.priority.value),
                ("Difficulty", rec.difficulty.value),
                ("Est. Time", rec.estimated_time),
            ]):
                cell = pdt_tbl.rows[0].cells[c_idx]
                cell.text = ""
                p1 = cell.paragraphs[0]
                r1 = p1.add_run(label)
                r1.bold = True
                r1.font.size = Pt(8)
                p2 = cell.add_paragraph(val)
                p2.runs[0].font.size = Pt(9)
                _set_cell_bg(cell, C_LIGHT_BG)

            # SonarQube link
            if rec.sonar_url:
                link_para = doc.add_paragraph()
                run = link_para.add_run(f"Open in SonarQube: {rec.sonar_url}")
                run.font.color.rgb = C_ACCENT
                run.font.size = Pt(8)

            _add_horizontal_rule(doc)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Top Files
    # ------------------------------------------------------------------

    def _top_files(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "6. Top Files")
        doc.add_paragraph("Top 20 files with the most quality issues.")

        headers = ["#", "File", "Total", "Bugs", "Code Smells", "Vulns"]
        tbl = doc.add_table(rows=1 + len(data.top_files), cols=len(headers))
        tbl.style = "Table Grid"

        for c_idx, h in enumerate(headers):
            cell = tbl.rows[0].cells[c_idx]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
            _set_cell_bg(cell, C_PRIMARY)

        for r_idx, fa in enumerate(data.top_files, start=1):
            row = tbl.rows[r_idx]
            vals = [str(r_idx), self.truncate(fa.path, 55),
                    str(fa.total_issues), str(fa.bugs),
                    str(fa.code_smells), str(fa.vulnerabilities)]
            for c_idx, val in enumerate(vals):
                cell = row.cells[c_idx]
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(8)
                if r_idx % 2 == 0:
                    _set_cell_bg(cell, C_LIGHT_BG)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Top Rules
    # ------------------------------------------------------------------

    def _top_rules(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "7. Top Rules")
        doc.add_paragraph("Top 20 most frequently violated rules.")

        headers = ["#", "Rule", "Violations", "Severity", "Type"]
        tbl = doc.add_table(rows=1 + len(data.top_rules), cols=len(headers))
        tbl.style = "Table Grid"

        for c_idx, h in enumerate(headers):
            cell = tbl.rows[0].cells[c_idx]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = C_WHITE
            _set_cell_bg(cell, C_PRIMARY)

        for r_idx, ra in enumerate(data.top_rules, start=1):
            row = tbl.rows[r_idx]
            vals = [str(r_idx), ra.rule_key, str(ra.count),
                    ra.severity, ra.issue_type.replace("_", " ").title()]
            for c_idx, val in enumerate(vals):
                cell = row.cells[c_idx]
                cell.text = val
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(8)
                if c_idx == 3:
                    color = SEV_COLORS.get(ra.severity)
                    if color:
                        run.bold = True
                        run.font.color.rgb = C_WHITE
                        _set_cell_bg(cell, color)
                elif r_idx % 2 == 0:
                    _set_cell_bg(cell, C_LIGHT_BG)

        doc.add_page_break()

    # ------------------------------------------------------------------
    # Technical Debt
    # ------------------------------------------------------------------

    def _technical_debt(self, doc: Document, data: ReportData) -> None:
        self._section_heading(doc, "8. Technical Debt")
        m = data.metrics

        doc.add_paragraph(
            f"The project carries a total technical debt of {m.technical_debt_display} "
            f"(debt ratio: {m.sqale_debt_ratio:.2f}%). "
            f"The estimated total remediation effort for all {len(data.issues)} identified issues "
            f"is approximately {data.total_effort_display}."
        )

        rows_data = [
            ["Total Technical Debt", m.technical_debt_display],
            ["Debt Ratio", f"{m.sqale_debt_ratio:.2f}%"],
            ["Total Estimated Effort", data.total_effort_display],
            ["Files Analysed", str(m.files)],
            ["Cyclomatic Complexity", str(m.complexity)],
            ["Cognitive Complexity", str(m.cognitive_complexity)],
            ["Comment Density", f"{m.comment_lines_density:.1f}%"],
        ]

        tbl = doc.add_table(rows=len(rows_data), cols=2)
        tbl.style = "Table Grid"
        for r_idx, (label, val) in enumerate(rows_data):
            row = tbl.rows[r_idx]
            row.cells[0].text = label
            row.cells[1].text = val
            row.cells[0].paragraphs[0].runs[0].bold = True
            if r_idx % 2 == 0:
                _set_cell_bg(row.cells[0], C_LIGHT_BG)
                _set_cell_bg(row.cells[1], C_LIGHT_BG)

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _section_heading(self, doc: Document, text: str) -> None:
        h = doc.add_heading(text, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h.runs:
            run.font.color.rgb = C_WHITE
        # Add background via shading
        pPr = h._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1A237E")
        pPr.append(shd)
        _set_para_spacing(h, before=14, after=6)

    def _sub_heading(self, doc: Document, text: str) -> None:
        para = doc.add_paragraph()
        run  = para.add_run(text)
        run.bold = True
        run.font.color.rgb = C_SECONDARY
        run.font.size = Pt(9)
        _set_para_spacing(para, before=6, after=2)

    @staticmethod
    def _strip_html(html: str) -> str:
        import re
        text = re.sub(r"<[^>]+>", " ", html or "")
        return re.sub(r"\s+", " ", text).strip()
